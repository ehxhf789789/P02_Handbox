"""
CNT 평가 시스템 FastAPI 서버
"""

import os
import sys
import json
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any, Optional

# 프로젝트 루트 추가
sys.path.insert(0, str(Path(__file__).parent.parent.parent))

from fastapi import FastAPI, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import boto3
from dotenv import load_dotenv

# 환경 변수 로드
load_dotenv()

from aws_agent.config import AWSConfig
from aws_agent.api.models.schemas import (
    EvaluateRequest, EvaluateResponse, EvidenceItem,
    FullEvaluateRequest, FullEvaluateResponse,
    ChunkRequest, ChunkResponse,
    EmbeddingRequest, EmbeddingResponse,
    SearchRequest, SearchResponse, SearchResult,
    LLMRequest, LLMResponse,
    MultiEvaluatorRequest, MultiEvaluatorResponse, EvaluatorResult,
    HealthResponse, EvaluatorType
)

# FastAPI 앱 생성
app = FastAPI(
    title="CNT 평가 시스템 API",
    description="건설신기술 평가를 위한 LLM 기반 API",
    version="1.0.0"
)

# CORS 설정
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 전역 설정
config = AWSConfig()


def get_bedrock_client():
    """Bedrock 클라이언트 생성"""
    return boto3.client(
        "bedrock-runtime",
        region_name=config.bedrock_region,
        aws_access_key_id=os.getenv("AWS_ACCESS_KEY_ID"),
        aws_secret_access_key=os.getenv("AWS_SECRET_ACCESS_KEY")
    )


# ===== 헬스 체크 =====

@app.get("/api/health", response_model=HealthResponse)
async def health_check():
    """헬스 체크"""
    aws_configured = bool(os.getenv("AWS_ACCESS_KEY_ID") and os.getenv("AWS_SECRET_ACCESS_KEY"))

    # Bedrock 연결 테스트
    bedrock_available = False
    if aws_configured:
        try:
            client = get_bedrock_client()
            bedrock_available = True
        except Exception:
            pass

    return HealthResponse(
        status="healthy",
        aws_configured=aws_configured,
        bedrock_available=bedrock_available,
        opensearch_available=bool(os.getenv("OPENSEARCH_ENDPOINT")),
        version="1.0.0"
    )


# ===== LLM 직접 호출 =====

def invoke_model_with_format(client, model_id: str, prompt: str, system_prompt: str = None, max_tokens: int = 1024, temperature: float = 0.1):
    """모델 형식에 맞게 호출"""

    # Anthropic Claude 형식
    if "anthropic" in model_id.lower() or "claude" in model_id.lower():
        messages = [{"role": "user", "content": prompt}]
        request_body = {
            "anthropic_version": "bedrock-2023-05-31",
            "max_tokens": max_tokens,
            "temperature": temperature,
            "messages": messages
        }
        if system_prompt:
            request_body["system"] = system_prompt

        response = client.invoke_model(
            modelId=model_id,
            body=json.dumps(request_body),
            contentType="application/json",
            accept="application/json"
        )
        response_body = json.loads(response["body"].read())
        return response_body["content"][0]["text"], response_body.get("usage")

    # Mistral 형식
    elif "mistral" in model_id.lower():
        full_prompt = prompt
        if system_prompt:
            full_prompt = f"<s>[INST] {system_prompt}\n\n{prompt} [/INST]"
        else:
            full_prompt = f"<s>[INST] {prompt} [/INST]"

        request_body = {
            "prompt": full_prompt,
            "max_tokens": max_tokens,
            "temperature": temperature,
        }

        response = client.invoke_model(
            modelId=model_id,
            body=json.dumps(request_body),
            contentType="application/json",
            accept="application/json"
        )
        response_body = json.loads(response["body"].read())
        output_text = response_body.get("outputs", [{}])[0].get("text", "")
        return output_text, None

    # Meta Llama 형식
    elif "llama" in model_id.lower() or "meta" in model_id.lower():
        if system_prompt:
            full_prompt = f"<|begin_of_text|><|start_header_id|>system<|end_header_id|>\n\n{system_prompt}<|eot_id|><|start_header_id|>user<|end_header_id|>\n\n{prompt}<|eot_id|><|start_header_id|>assistant<|end_header_id|>\n\n"
        else:
            full_prompt = f"<|begin_of_text|><|start_header_id|>user<|end_header_id|>\n\n{prompt}<|eot_id|><|start_header_id|>assistant<|end_header_id|>\n\n"

        request_body = {
            "prompt": full_prompt,
            "max_gen_len": max_tokens,
            "temperature": temperature,
        }

        response = client.invoke_model(
            modelId=model_id,
            body=json.dumps(request_body),
            contentType="application/json",
            accept="application/json"
        )
        response_body = json.loads(response["body"].read())
        return response_body.get("generation", ""), None

    # Amazon Titan 형식
    else:
        request_body = {
            "inputText": f"{system_prompt}\n\n{prompt}" if system_prompt else prompt,
            "textGenerationConfig": {
                "maxTokenCount": max_tokens,
                "temperature": temperature,
            }
        }

        response = client.invoke_model(
            modelId=model_id,
            body=json.dumps(request_body),
            contentType="application/json",
            accept="application/json"
        )
        response_body = json.loads(response["body"].read())
        return response_body.get("results", [{}])[0].get("outputText", ""), None


@app.post("/api/llm/invoke", response_model=LLMResponse)
async def invoke_llm(request: LLMRequest):
    """Bedrock LLM 직접 호출 (Claude, Mistral, Llama 지원)"""
    try:
        client = get_bedrock_client()

        response_text, usage = invoke_model_with_format(
            client=client,
            model_id=config.bedrock_model_id,
            prompt=request.prompt,
            system_prompt=request.system_prompt,
            max_tokens=request.max_tokens,
            temperature=request.temperature
        )

        return LLMResponse(
            success=True,
            response=response_text,
            model=config.bedrock_model_id,
            usage=usage
        )

    except Exception as e:
        return LLMResponse(
            success=False,
            response="",
            model=config.bedrock_model_id,
            error=str(e)
        )


# ===== 단일 평가 =====

@app.post("/api/evaluate", response_model=EvaluateResponse)
async def evaluate_single(request: EvaluateRequest):
    """단일 평가 실행"""
    try:
        # 평가 유형에 따른 에이전트 선택
        if request.evaluator_type == EvaluatorType.NOVELTY:
            from aws_agent.agents.novelty_agent import NoveltyEvaluationAgent
            agent = NoveltyEvaluationAgent(config)
        elif request.evaluator_type == EvaluatorType.PROGRESSIVENESS:
            from aws_agent.agents.progress_agent import ProgressEvaluationAgent
            agent = ProgressEvaluationAgent(config)
        elif request.evaluator_type == EvaluatorType.FIELD:
            from aws_agent.agents.field_agent import FieldExcellenceAgent
            agent = FieldExcellenceAgent(config)
        else:
            raise ValueError(f"Unknown evaluator type: {request.evaluator_type}")

        # 평가 실행
        result = agent.evaluate(request.tech_id)

        return EvaluateResponse(
            success=True,
            tech_id=request.tech_id,
            evaluator_type=request.evaluator_type.value,
            score=result.score,
            grade=result.grade,
            pass_status=result.pass_status,
            evidence=[EvidenceItem(**e) for e in result.evidence] if result.evidence else [],
            comments=result.comments,
            sub_scores=result.sub_scores
        )

    except Exception as e:
        return EvaluateResponse(
            success=False,
            tech_id=request.tech_id,
            evaluator_type=request.evaluator_type.value,
            score=0,
            grade="오류",
            pass_status=False,
            error=str(e)
        )


# ===== 전체 평가 =====

@app.post("/api/evaluate/full", response_model=FullEvaluateResponse)
async def evaluate_full(request: FullEvaluateRequest):
    """전체 평가 실행 (1차 + 2차 심사)"""
    try:
        from aws_agent.evaluation.runner import EvaluationRunner

        runner = EvaluationRunner(config, use_aws=request.use_aws)
        result = runner.run_full_evaluation(request.tech_id)

        return FullEvaluateResponse(
            success=True,
            tech_id=result.tech_id,
            evaluation_date=result.evaluation_date,
            stage_1=result.stage_1,
            stage_2=result.stage_2,
            overall_score=result.overall_score,
            overall_pass=result.overall_pass,
            summary=result.summary
        )

    except Exception as e:
        return FullEvaluateResponse(
            success=False,
            tech_id=request.tech_id,
            evaluation_date=datetime.now().isoformat(),
            stage_1={},
            stage_2={},
            overall_score=0,
            overall_pass=False,
            summary={},
            error=str(e)
        )


# ===== 10명 평가위원 다중 평가 =====

EVALUATOR_CONFIGS = [
    {"id": "1", "expertise": "structure", "stance": "conservative", "eval_type": "novelty"},
    {"id": "2", "expertise": "construction", "stance": "progressive", "eval_type": "novelty"},
    {"id": "3", "expertise": "materials", "stance": "neutral", "eval_type": "progressiveness"},
    {"id": "4", "expertise": "economics", "stance": "neutral", "eval_type": "progressiveness"},
    {"id": "5", "expertise": "patent", "stance": "conservative", "eval_type": "novelty"},
    {"id": "6", "expertise": "safety", "stance": "conservative", "eval_type": "field"},
    {"id": "7", "expertise": "environmental", "stance": "progressive", "eval_type": "field"},
    {"id": "8", "expertise": "geotechnical", "stance": "neutral", "eval_type": "progressiveness"},
    {"id": "9", "expertise": "policy", "stance": "neutral", "eval_type": "field"},
    {"id": "10", "expertise": "sustainability", "stance": "progressive", "eval_type": "novelty"},
]


@app.post("/api/evaluate/multi", response_model=MultiEvaluatorResponse)
async def evaluate_multi(request: MultiEvaluatorRequest):
    """10명 평가위원 다중 평가"""
    try:
        client = get_bedrock_client()
        evaluator_results = []

        # 평가위원 설정
        evaluators = request.evaluators if request.evaluators else EVALUATOR_CONFIGS

        for eval_config in evaluators:
            evaluator_id = eval_config.get("id", "unknown")
            expertise = eval_config.get("expertise", "general")
            stance = eval_config.get("stance", "neutral")
            eval_type = eval_config.get("eval_type", "novelty")

            # 평가위원별 시스템 프롬프트
            system_prompt = f"""당신은 건설신기술 심사위원회의 {expertise} 분야 전문가입니다.
평가 성향: {stance}
평가 관점: {"보수적이고 안전성을 중시" if stance == "conservative" else "혁신적이고 기술 발전을 중시" if stance == "progressive" else "균형 잡힌 관점"}

평가 항목:
1. 신규성 (50점): 기존기술과의 차별성 (25점), 독창성과 자립성 (25점)
2. 진보성 (50점): 품질 향상 (15점), 개발 정도 (15점), 안전성 (10점), 첨단기술성 (10점)

각 항목별로 점수와 구체적 근거를 제시하고, 최종 판정(통과/불통과)을 결정하세요.

반드시 다음 JSON 형식으로 응답하세요:
```json
{{
  "verdict": "Approved" 또는 "Rejected",
  "novelty_score": 0-50 사이 점수,
  "progress_score": 0-50 사이 점수,
  "confidence": 0.0-1.0 사이 신뢰도,
  "evidence": [
    {{"type": "신규성/진보성", "content": "근거 내용", "location": "문서 위치"}}
  ],
  "comments": "종합 평가 의견"
}}
```"""

            # 평가 프롬프트
            eval_prompt = f"""신기술 번호: {request.tech_id}

문서 컨텍스트:
{request.document_context or "문서 정보가 제공되지 않았습니다."}

위 건설신기술에 대해 {expertise} 분야 전문가로서 평가해주세요."""

            # Bedrock 호출 (다양한 모델 형식 지원)
            response_text, _ = invoke_model_with_format(
                client=client,
                model_id=config.bedrock_model_id,
                prompt=eval_prompt,
                system_prompt=system_prompt,
                max_tokens=2048,
                temperature=0.2
            )

            # JSON 파싱
            try:
                import re
                json_match = re.search(r'```json\s*(.*?)\s*```', response_text, re.DOTALL)
                if json_match:
                    eval_data = json.loads(json_match.group(1))
                else:
                    json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
                    eval_data = json.loads(json_match.group(0)) if json_match else {}
            except:
                eval_data = {
                    "verdict": "Rejected",
                    "novelty_score": 30,
                    "progress_score": 30,
                    "confidence": 0.5,
                    "evidence": [],
                    "comments": "응답 파싱 실패"
                }

            evaluator_results.append(EvaluatorResult(
                evaluator_id=f"evaluator_{evaluator_id}",
                expertise=expertise,
                stance=stance,
                verdict=eval_data.get("verdict", "Rejected"),
                novelty_score=eval_data.get("novelty_score", 30),
                progress_score=eval_data.get("progress_score", 30),
                confidence=eval_data.get("confidence", 0.7),
                citation=eval_data.get("comments", ""),
                evidence=eval_data.get("evidence", [])
            ))

        # 집계
        approved_count = sum(1 for r in evaluator_results if r.verdict == "Approved")
        rejected_count = len(evaluator_results) - approved_count
        final_verdict = "Approved" if approved_count >= 6 else "Rejected"

        return MultiEvaluatorResponse(
            success=True,
            tech_id=request.tech_id,
            evaluator_results=evaluator_results,
            approved_count=approved_count,
            rejected_count=rejected_count,
            final_verdict=final_verdict
        )

    except Exception as e:
        return MultiEvaluatorResponse(
            success=False,
            tech_id=request.tech_id,
            evaluator_results=[],
            approved_count=0,
            rejected_count=0,
            final_verdict="Error",
            error=str(e)
        )


# ===== 문서 청킹 =====

@app.post("/api/chunk", response_model=ChunkResponse)
async def chunk_document(request: ChunkRequest):
    """텍스트 청킹"""
    try:
        chunks = []
        text = request.text
        chunk_size = request.chunk_size
        overlap = request.chunk_overlap

        # 슬라이딩 윈도우 청킹
        for i in range(0, len(text), chunk_size - overlap):
            chunk_text = text[i:i + chunk_size]
            if chunk_text.strip():
                chunks.append({
                    "index": len(chunks),
                    "content": chunk_text,
                    "start_pos": i,
                    "end_pos": min(i + chunk_size, len(text))
                })

        return ChunkResponse(
            success=True,
            chunks=chunks,
            chunks_count=len(chunks)
        )

    except Exception as e:
        return ChunkResponse(
            success=False,
            chunks=[],
            chunks_count=0,
            error=str(e)
        )


# ===== 임베딩 생성 =====

@app.post("/api/embedding", response_model=EmbeddingResponse)
async def create_embeddings(request: EmbeddingRequest):
    """Bedrock Titan 임베딩 생성"""
    try:
        client = get_bedrock_client()
        embeddings = []

        for text in request.texts:
            # Titan Embeddings 호출
            response = client.invoke_model(
                modelId=config.bedrock_embedding_model_id,
                body=json.dumps({"inputText": text[:8000]}),  # 최대 8000자
                contentType="application/json",
                accept="application/json"
            )

            response_body = json.loads(response["body"].read())
            embedding = response_body.get("embedding", [])
            embeddings.append(embedding)

        return EmbeddingResponse(
            success=True,
            embeddings=embeddings,
            dimension=len(embeddings[0]) if embeddings else 0
        )

    except Exception as e:
        return EmbeddingResponse(
            success=False,
            embeddings=[],
            dimension=0,
            error=str(e)
        )


# ===== 검색 (OpenSearch 또는 로컬 벡터스토어) =====

@app.post("/api/search", response_model=SearchResponse)
async def search_documents(request: SearchRequest):
    """문서 검색 (OpenSearch 또는 로컬 벡터스토어 - 시뮬레이션 없음)"""
    try:
        from aws_agent.preprocessing.embedder import BedrockEmbedder

        # Bedrock 임베딩으로 쿼리 벡터 생성 (API 필수)
        embedder = BedrockEmbedder(config)
        query_embedding = embedder.embed_query(request.query)

        # OpenSearch가 설정된 경우
        if os.getenv("OPENSEARCH_ENDPOINT"):
            from aws_agent.vectorstore.opensearch_client import OpenSearchVectorStore
            vectorstore = OpenSearchVectorStore(config)
        else:
            # 로컬 벡터스토어 사용 (ChromaDB)
            from aws_agent.vectorstore.opensearch_client import LocalVectorStore
            vectorstore = LocalVectorStore()

        results = vectorstore.search(
            query_embedding=query_embedding,
            tech_id=request.tech_id,
            k=request.k
        )

        return SearchResponse(
            success=True,
            results=[
                SearchResult(
                    content=r.get("content", ""),
                    section=r.get("section"),
                    page_numbers=r.get("page_numbers", []),
                    score=r.get("_score", r.get("score", 0))
                )
                for r in results
            ],
            total=len(results)
        )

    except Exception as e:
        return SearchResponse(
            success=False,
            results=[],
            total=0,
            error=f"검색 실패 (Bedrock API 연결 필요): {str(e)}"
        )


# ===== 문서 파싱 (PDF/HWP/XLSX) =====

class ParseRequest(BaseModel):
    file_path: str

class ParseResponse(BaseModel):
    success: bool
    text: str = ""
    characters: int = 0
    pages: int = 0
    file_path: str = ""
    error: Optional[str] = None

@app.post("/api/parse", response_model=ParseResponse)
async def parse_document(request: ParseRequest):
    """문서 파싱 (PDF/HWP/XLSX/DOCX)"""
    try:
        from pathlib import Path
        file_path = Path(request.file_path)

        if not file_path.exists():
            return ParseResponse(
                success=False,
                file_path=request.file_path,
                error="파일을 찾을 수 없습니다"
            )

        ext = file_path.suffix.lower()
        text = ""
        pages = 0

        if ext == ".pdf":
            # PyMuPDF로 PDF 파싱
            import fitz
            doc = fitz.open(str(file_path))
            pages = len(doc)
            text_parts = []
            for page in doc:
                text_parts.append(page.get_text())
            doc.close()
            text = "\n".join(text_parts)

        elif ext == ".xlsx" or ext == ".xls":
            # openpyxl로 Excel 파싱
            import openpyxl
            wb = openpyxl.load_workbook(str(file_path), data_only=True)
            text_parts = []
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                text_parts.append(f"\n[시트: {sheet}]\n")
                for row in ws.iter_rows(values_only=True):
                    row_text = " | ".join(str(cell) if cell else "" for cell in row)
                    if row_text.strip():
                        text_parts.append(row_text)
            text = "\n".join(text_parts)
            pages = len(wb.sheetnames)

        elif ext == ".docx":
            # python-docx로 Word 파싱
            from docx import Document
            doc = Document(str(file_path))
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            pages = len(doc.paragraphs) // 30 + 1

        elif ext == ".hwp":
            # olefile로 HWP 파싱 (기본 텍스트 추출)
            import olefile
            if olefile.isOleFile(str(file_path)):
                ole = olefile.OleFileIO(str(file_path))
                if ole.exists("PrvText"):
                    text = ole.openstream("PrvText").read().decode("utf-16", errors="ignore")
                ole.close()
            else:
                text = "[HWP 파싱 실패: OLE 형식이 아님]"
            pages = 1

        elif ext == ".txt":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            pages = 1

        else:
            return ParseResponse(
                success=False,
                file_path=request.file_path,
                error=f"지원하지 않는 파일 형식: {ext}"
            )

        return ParseResponse(
            success=True,
            text=text,
            characters=len(text),
            pages=pages,
            file_path=request.file_path
        )

    except Exception as e:
        return ParseResponse(
            success=False,
            file_path=request.file_path,
            error=str(e)
        )


# ===== OpenSearch 인덱싱 =====

class IndexRequest(BaseModel):
    tech_id: str
    chunks: List[Dict[str, Any]]

class IndexResponse(BaseModel):
    success: bool
    indexed_count: int = 0
    collection: str = ""
    error: Optional[str] = None

@app.post("/api/index", response_model=IndexResponse)
async def index_documents(request: IndexRequest):
    """문서를 OpenSearch에 인덱싱"""
    try:
        from aws_agent.vectorstore.opensearch_client import OpenSearchVectorStore, LocalVectorStore
        from aws_agent.preprocessing.embedder import BedrockEmbedder

        # 임베딩 생성
        embedder = BedrockEmbedder(config)
        embedded_docs = []

        for chunk in request.chunks:
            content = chunk.get("content", "")
            if not content:
                continue

            embedding = embedder.create_embedding(content)

            embedded_docs.append({
                "tech_id": request.tech_id,
                "chunk_index": chunk.get("index", len(embedded_docs)),
                "content": content,
                "embedding": embedding,
                "page_numbers": chunk.get("page_numbers", []),
                "section": chunk.get("section", ""),
            })

        # OpenSearch 또는 로컬 저장소에 인덱싱
        if os.getenv("OPENSEARCH_ENDPOINT"):
            vectorstore = OpenSearchVectorStore(config)
            vectorstore.create_index()
            indexed = vectorstore.index_documents(embedded_docs)
        else:
            # 로컬 벡터스토어 사용
            vectorstore = LocalVectorStore()
            indexed = vectorstore.index_documents(embedded_docs)

        return IndexResponse(
            success=True,
            indexed_count=indexed,
            collection=config.opensearch_index_name or "local"
        )

    except Exception as e:
        return IndexResponse(
            success=False,
            error=str(e)
        )


# ===== 전체 RAG 파이프라인 =====

class RAGPipelineRequest(BaseModel):
    tech_id: str
    file_paths: List[str]
    evaluate: bool = True

class RAGPipelineResponse(BaseModel):
    success: bool
    tech_id: str
    files_parsed: int = 0
    chunks_created: int = 0
    vectors_indexed: int = 0
    evaluation_result: Optional[Dict[str, Any]] = None
    error: Optional[str] = None

@app.post("/api/pipeline", response_model=RAGPipelineResponse)
async def run_rag_pipeline(request: RAGPipelineRequest):
    """전체 RAG 파이프라인: 파싱 → 청킹 → 임베딩 → 인덱싱 → 평가"""
    try:
        from pathlib import Path
        import fitz  # PyMuPDF

        all_text = []
        files_parsed = 0

        # 1. 문서 파싱
        for file_path in request.file_paths:
            path = Path(file_path)
            if not path.exists():
                continue

            ext = path.suffix.lower()
            text = ""

            if ext == ".pdf":
                doc = fitz.open(str(path))
                for page in doc:
                    text += page.get_text() + "\n"
                doc.close()
            elif ext == ".txt":
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
            else:
                continue

            all_text.append(text)
            files_parsed += 1

        combined_text = "\n\n".join(all_text)

        # 2. 청킹
        chunks = []
        chunk_size = config.chunk_size
        overlap = config.chunk_overlap

        for i in range(0, len(combined_text), chunk_size - overlap):
            chunk_text = combined_text[i:i + chunk_size]
            if chunk_text.strip():
                chunks.append({
                    "index": len(chunks),
                    "content": chunk_text,
                })

        # 3. 임베딩 + 인덱싱
        from aws_agent.preprocessing.embedder import BedrockEmbedder
        from aws_agent.vectorstore.opensearch_client import OpenSearchVectorStore, LocalVectorStore

        embedder = BedrockEmbedder(config)
        embedded_docs = []

        for chunk in chunks:
            embedding = embedder.create_embedding(chunk["content"])
            embedded_docs.append({
                "tech_id": request.tech_id,
                "chunk_index": chunk["index"],
                "content": chunk["content"],
                "embedding": embedding,
            })

        # 인덱싱
        if os.getenv("OPENSEARCH_ENDPOINT"):
            vectorstore = OpenSearchVectorStore(config)
            vectorstore.create_index()
            indexed = vectorstore.index_documents(embedded_docs)
        else:
            vectorstore = LocalVectorStore()
            indexed = vectorstore.index_documents(embedded_docs)

        # 4. 평가 (선택적)
        evaluation_result = None
        if request.evaluate:
            from aws_agent.evaluation.runner import EvaluationRunner
            runner = EvaluationRunner(config, use_aws=True)
            result = runner.run_full_evaluation(request.tech_id)
            evaluation_result = result.to_dict()

        return RAGPipelineResponse(
            success=True,
            tech_id=request.tech_id,
            files_parsed=files_parsed,
            chunks_created=len(chunks),
            vectors_indexed=indexed,
            evaluation_result=evaluation_result
        )

    except Exception as e:
        return RAGPipelineResponse(
            success=False,
            tech_id=request.tech_id,
            error=str(e)
        )


# ===== KISTI ScienceON API =====

class KISTISearchRequest(BaseModel):
    """KISTI 검색 요청"""
    target: str  # ARTI, PATENT, REPORT, ATT
    query: str
    search_field: str = "BI"  # BI: 기본색인, TI: 제목, AU: 저자
    cur_page: int = 1
    row_count: int = 10

class KISTISearchResponse(BaseModel):
    """KISTI 검색 응답"""
    success: bool
    target: str = ""
    total_count: int = 0
    current_page: int = 1
    records_count: int = 0
    records: List[Dict[str, Any]] = []
    error: Optional[str] = None

class KISTITokenResponse(BaseModel):
    """KISTI 토큰 상태 응답"""
    success: bool
    has_valid_token: bool = False
    access_token_expire: Optional[str] = None
    refresh_token_expire: Optional[str] = None
    error: Optional[str] = None

# KISTI 클라이언트 인스턴스 (지연 로딩)
_kisti_client = None

def get_kisti_client():
    """KISTI 클라이언트 싱글톤"""
    global _kisti_client
    if _kisti_client is None:
        try:
            from aws_agent.api.kisti_client import KISTIClient
            _kisti_client = KISTIClient()
        except ImportError as e:
            print(f"[경고] KISTI 클라이언트 로드 실패: {e}")
            print("pip install pycryptodome 실행 필요")
            return None
    return _kisti_client


@app.get("/api/kisti/token", response_model=KISTITokenResponse)
async def kisti_token_status():
    """KISTI 토큰 상태 확인 및 발급"""
    client = get_kisti_client()
    if client is None:
        return KISTITokenResponse(
            success=False,
            error="KISTI 클라이언트 초기화 실패 (pycryptodome 설치 필요)"
        )

    try:
        # 토큰 발급/갱신
        if not client.ensure_valid_token():
            return KISTITokenResponse(
                success=False,
                error="토큰 발급 실패"
            )

        return KISTITokenResponse(
            success=True,
            has_valid_token=client.token_info.is_access_token_valid(),
            access_token_expire=client.token_info.access_token_expire.isoformat() if client.token_info.access_token_expire else None,
            refresh_token_expire=client.token_info.refresh_token_expire.isoformat() if client.token_info.refresh_token_expire else None
        )
    except Exception as e:
        return KISTITokenResponse(
            success=False,
            error=str(e)
        )


@app.post("/api/kisti/search", response_model=KISTISearchResponse)
async def kisti_search(request: KISTISearchRequest):
    """
    KISTI ScienceON 검색 API

    서비스 타입 (target):
    - ARTI: 논문
    - PATENT: 특허
    - REPORT: 보고서
    - ATT: 동향

    검색 필드 (search_field):
    - BI: 기본색인 (기본값)
    - TI: 제목
    - AU: 저자
    - AB: 초록
    - KW: 키워드
    """
    client = get_kisti_client()
    if client is None:
        return KISTISearchResponse(
            success=False,
            error="KISTI 클라이언트 초기화 실패 (pycryptodome 설치 필요)"
        )

    try:
        result = client.search(
            target=request.target,
            query=request.query,
            search_field=request.search_field,
            cur_page=request.cur_page,
            row_count=request.row_count
        )

        if result.get("success"):
            return KISTISearchResponse(
                success=True,
                target=result.get("target", request.target),
                total_count=result.get("total_count", 0),
                current_page=result.get("current_page", 1),
                records_count=result.get("records_count", 0),
                records=result.get("records", [])
            )
        else:
            return KISTISearchResponse(
                success=False,
                target=request.target,
                error=result.get("error", "검색 실패")
            )

    except Exception as e:
        return KISTISearchResponse(
            success=False,
            target=request.target,
            error=str(e)
        )


@app.get("/api/kisti/services")
async def kisti_services():
    """KISTI 지원 서비스 목록"""
    return {
        "services": [
            {"code": "ARTI", "name": "논문", "description": "국내외 학술논문 검색"},
            {"code": "PATENT", "name": "특허", "description": "국내외 특허정보 검색"},
            {"code": "REPORT", "name": "보고서", "description": "연구보고서 검색"},
            {"code": "ATT", "name": "동향", "description": "과학기술 동향정보 검색"}
        ],
        "search_fields": [
            {"code": "BI", "name": "기본색인", "description": "모든 필드에서 검색"},
            {"code": "TI", "name": "제목", "description": "제목에서만 검색"},
            {"code": "AU", "name": "저자", "description": "저자명에서만 검색"},
            {"code": "AB", "name": "초록", "description": "초록에서만 검색"},
            {"code": "KW", "name": "키워드", "description": "키워드에서만 검색"}
        ],
        "search_operators": [
            {"operator": "()", "description": "우선순위 지정"},
            {"operator": "공백", "description": "AND 연산 (모두 포함)"},
            {"operator": "|", "description": "OR 연산 (하나 이상 포함)"},
            {"operator": "!", "description": "NOT 연산 (제외)"},
            {"operator": "*", "description": "와일드카드 (0개 이상 문자)"},
            {"operator": '""', "description": "구문 일치 검색"}
        ]
    }


# ===== 메인 =====

if __name__ == "__main__":
    import uvicorn
    print("=" * 60)
    print("CNT 평가 시스템 API 서버")
    print(f"AWS Region: {config.bedrock_region}")
    print(f"Bedrock Model: {config.bedrock_model_id}")
    print(f"Embedding Model: {config.bedrock_embedding_model_id}")
    print(f"OpenSearch: {os.getenv('OPENSEARCH_ENDPOINT', 'Not configured (using local)')}")
    print("KISTI ScienceON API: 활성화")
    print("=" * 60)
    uvicorn.run(app, host="127.0.0.1", port=8000)
